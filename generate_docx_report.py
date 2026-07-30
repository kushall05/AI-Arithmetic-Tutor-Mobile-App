import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def build_college_report():
    doc = Document()

    # Page Margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Font Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.25
    style_normal.paragraph_format.space_after = Pt(6)

    # Base Image Paths
    img_dir = r"C:\Users\saras\.gemini\antigravity\brain\a7fe2f1e-46fe-4667-a3f9-7194962c1929"
    dash_img = os.path.join(img_dir, "dashboard_screen_1785445190282.jpg")
    prac_img = os.path.join(img_dir, "practice_screen_1785445203275.jpg")
    tutor_img = os.path.join(img_dir, "tutor_screen_1785445216888.jpg")
    prog_img = os.path.join(img_dir, "progress_screen_1785445291237.jpg")
    prof_img = os.path.join(img_dir, "profile_screen_1785445323655.jpg")

    # Header Helper
    def add_header(p_text_left="STUDENT NAME: KUSHALL THORAT", p_text_right="ROLL NO - 39038"):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell_l, cell_r = table.rows[0].cells
        cell_l.width = Inches(3.25)
        cell_r.width = Inches(3.25)

        p_l = cell_l.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_l = p_l.add_run(p_text_left)
        run_l.bold = True
        run_l.font.size = Pt(11)

        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_r = p_r.add_run(p_text_right)
        run_r.bold = True
        run_r.font.size = Pt(11)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- TOP HEADER ---
    add_header()

    # --- INDEX / TABLE OF CONTENTS ---
    h_idx = doc.add_paragraph()
    r_idx = h_idx.add_run("INDEX")
    r_idx.bold = True
    r_idx.font.size = Pt(16)
    r_idx.underline = True
    h_idx.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_idx.paragraph_format.space_after = Pt(12)

    # Index Table
    idx_table = doc.add_table(rows=7, cols=3)
    idx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx_table.autofit = False

    col_widths = [Inches(0.8), Inches(4.7), Inches(1.0)]
    headers = ["SR NO", "TOPIC NAME", "PAGE NO"]

    hdr_cells = idx_table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(11)

    topics = [
        ("1", "PROJECT TITLE", "1"),
        ("2", "PROJECT DESCRIPTION", "1"),
        ("3", "SYSTEM WORKFLOW / INTEGRATION FLOWCHART", "2"),
        ("4", "HOW ALL COMPONENTS ARE INTEGRATED", "3"),
        ("5", "WEBSITE & MOBILE APP SCREENSHOTS", "5"),
        ("6", "OUTPUT & DEMONSTRATION", "8")
    ]

    for row_idx, (sr, title, pg) in enumerate(topics, start=1):
        cells = idx_table.rows[row_idx].cells
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
        cells[2].width = col_widths[2]

        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.add_run(sr)

        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.add_run(title)

        p2 = cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(pg)

    doc.add_page_break()

    # --- 1. PROJECT TITLE ---
    add_header()
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. PROJECT TITLE:")
    r1.bold = True
    r1.font.size = Pt(14)
    r1.underline = True

    p_title = doc.add_paragraph()
    r_tval = p_title.add_run("AI ARITHMETIC TUTOR – AN INTELLIGENT MOBILE LEARNING APPLICATION")
    r_tval.bold = True
    r_tval.font.size = Pt(13)
    p_title.paragraph_format.space_after = Pt(14)

    # --- 2. PROJECT DESCRIPTION ---
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. PROJECT DESCRIPTION:")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.underline = True

    desc_p1 = doc.add_paragraph()
    desc_p1.add_run("The ").font.size = Pt(12)
    r_bold_name = desc_p1.add_run("AI Arithmetic Tutor")
    r_bold_name.bold = True
    desc_p1.add_run(" is an Artificial Intelligence-based hybrid mobile and web application developed to assist elementary and middle school students in mastering fundamental arithmetic skills (Addition, Subtraction, Multiplication, and Division). Arithmetic represents the foundational cornerstone of mathematical literacy. However, due to individual learning paces, students frequently experience procedural errors, regrouping confusion, and conceptual gaps.")

    desc_p2 = doc.add_paragraph()
    desc_p2.add_run("The objective of this project is to create an intelligent system that can dynamically generate practice questions, evaluate student answer submissions in real-time, explain mistake reasoning step-by-step using Google Gemini AI, and provide Socratic hints without directly revealing final numerical answers. The platform dynamically adjusts problem difficulty based on historical accuracy.")

    desc_p3 = doc.add_paragraph()
    desc_p3.add_run("The application allows users to register, practice arithmetic operations across multiple difficulty levels (Easy, Medium, Hard, Challenge), attempt timed quiz marathons, interact verbally using Web Speech synthesis and speech recognition, and draw scratch calculations on an integrated canvas scratchpad. The backend is developed using Python Flask with an embedded SQLite database (tutor.db) for managing user accounts, operation mastery, quiz history, AI tutor interaction logs, and unlocked achievement badges.")

    desc_p4 = doc.add_paragraph()
    desc_p4.add_run("Furthermore, the project includes an automated PDF report generator using ReportLab that synthesizes user performance statistics, accuracy percentages, mastery star levels, and AI recommendations into an exportable academic document. The application supports cross-platform hybrid deployment operating as a Web Application (Flask/Gunicorn), a Progressive Web App (PWA), and a Native Android WebView APK built automatically via GitHub Actions CI/CD.")

    doc.add_page_break()

    # --- 3. SYSTEM WORKFLOW / INTEGRATION FLOWCHART ---
    add_header()
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. SYSTEM WORKFLOW / INTEGRATION FLOWCHART:")
    r3.bold = True
    r3.font.size = Pt(14)
    r3.underline = True

    flowchart_text = """
                                  USER
                                   |
                                   v
             [ Select Operation & Solve Arithmetic Problem ]
                                   |
                                   v
             Frontend Mobile / PWA Interface (HTML5 + CSS3 + JS)
             - Web Speech Voice Input & Canvas Scratchpad
                                   |
                                   v
             Flask Backend Server REST API (app.py)
             - Auth & Session Tracking
             - Score & Streak Counter Update
                                   |
                                   v
             Database Storage (tutor.db SQLite)
             - Users, Progress, Quiz Results, Badges
                                   |
                                   v
             AI Tutor Engine (ai_engine.py)
             - Gemini API / Fallback Step-by-Step Solver
             - Socratic Hint Generator
                                   |
            +----------------------+----------------------+
            |                                             |
            v                                             v
[ Step-by-Step Mistake Breakdown ]              [ Adaptive Quiz & Score ]
            |                                             |
            +----------------------+----------------------+
                                   |
                                   v
             Progress Analytics Dashboard & PDF Report Export
    """

    p_flow = doc.add_paragraph()
    r_flow = p_flow.add_run(flowchart_text)
    r_flow.font.name = 'Courier New'
    r_flow.font.size = Pt(9.5)
    p_flow.paragraph_format.space_after = Pt(14)

    doc.add_page_break()

    # --- 4. HOW ALL COMPONENTS ARE INTEGRATED ---
    add_header()
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. HOW ALL COMPONENTS ARE INTEGRATED:")
    r4.bold = True
    r4.font.size = Pt(14)
    r4.underline = True

    p_int_intro = doc.add_paragraph()
    p_int_intro.add_run("The AI Arithmetic Tutor follows a structured software architecture where the mobile frontend, Flask REST backend, SQLite database, Gemini AI processing module, voice engine, and PDF generator work together to provide a complete interactive workflow.")

    # Subsection 1
    h4_1 = doc.add_paragraph()
    r4_1 = h4_1.add_run("1. Frontend and User Interaction Integration")
    r4_1.bold = True
    r4_1.font.size = Pt(12)

    p4_1 = doc.add_paragraph()
    p4_1.add_run("The frontend acts as the first layer where users interact with the application:\n"
               "• HTML5 for single-page application (SPA) screen structure.\n"
               "• Vanilla CSS3 for an eye-friendly cool dark navy theme, glassmorphic UI cards, and responsive smartphone shell layout.\n"
               "• JavaScript for app routing, practice problem evaluation, canvas scratchpad drawing, and Chart.js progress visualization.\n"
               "• Web Speech API for reading questions aloud (TTS) and receiving spoken numeric answers.")

    # Subsection 2
    h4_2 = doc.add_paragraph()
    r4_2 = h4_2.add_run("2. Frontend Integration with Flask Backend")
    r4_2.bold = True
    r4_2.font.size = Pt(12)

    p4_2 = doc.add_paragraph()
    p4_2.add_run("The Flask framework (app.py) serves as the communication bridge between the UI and AI module:\n"
               "• Manages REST endpoints (/api/auth/*, /api/practice/*, /api/quiz/*, /api/ai/*, /api/progress/*).\n"
               "• Validates answer submissions, updates practice streak counts, awards XP points (+10 XP per correct answer, +50 XP daily challenge), and checks badge criteria.\n"
               "• Interacts with database.py to maintain user state and statistics.")

    # Subsection 3
    h4_3 = doc.add_paragraph()
    r4_3 = h4_3.add_run("3. Flask Backend Integration with AI Detection & Tutor Module")
    r4_3.bold = True
    r4_3.font.size = Pt(12)

    p4_3 = doc.add_paragraph()
    p4_3.add_run("The AI module (ai_engine.py) connects directly to Google Gemini API:\n"
               "• Calls Gemini 1.5/2.0 API to generate encouraging 3-step error explanations.\n"
               "• Implements a Socratic Hint Generator offering 3 progressive clues without revealing the answer.\n"
               "• Includes a Deterministic Step-by-Step Fallback Solver to guarantee 100% offline uptime if API keys are absent.")

    # Subsection 4
    h4_4 = doc.add_paragraph()
    r4_4 = h4_4.add_run("4. Database & PDF Report Generator Integration")
    r4_4.bold = True
    r4_4.font.size = Pt(12)

    p4_4 = doc.add_paragraph()
    p4_4.add_run("• Database (tutor.db): Manages SQLite relational tables for Users, Questions, Quiz Results, Progress, AI History, and Badges.\n"
               "• PDF Generator (pdf_generator.py): Leverages ReportLab to synthesize user statistics, operation accuracy percentages, unlocked badges, and AI teacher recommendations into an exportable academic document.")

    doc.add_page_break()

    # --- 5. WEBSITE & MOBILE APP SCREENSHOTS ---
    add_header()
    h5 = doc.add_paragraph()
    r5 = h5.add_run("5. WEBSITE & MOBILE APP SCREENSHOTS:")
    r5.bold = True
    r5.font.size = Pt(14)
    r5.underline = True

    screenshots = [
        ("Figure 1: Home Dashboard Screen (Greeting, Daily Challenge Banner, Operations Grid, XP & Streak Pill)", dash_img),
        ("Figure 2: Addition Practice Screen (Question Card, Multiple Choice Grid, Voice & Hint Buttons)", prac_img),
        ("Figure 3: AI Tutor Chat Screen (Professor Owl Interactive Step-by-Step Explanation)", tutor_img),
        ("Figure 4: Learning Progress & Analytics Screen (Operation Mastery Bars & Accuracy Chart)", prog_img),
        ("Figure 5: Profile & Settings Screen (DemoStudent Avatar, Grade Selector, Gemini API Input, PDF Export)", prof_img)
    ]

    for title, img_path in screenshots:
        p_img_title = doc.add_paragraph()
        r_it = p_img_title.add_run(title)
        r_it.bold = True
        r_it.font.size = Pt(11)

        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(3.2))
            p_img.paragraph_format.space_after = Pt(14)
        else:
            p_missing = doc.add_paragraph(f"[Image file not found: {img_path}]")
            p_missing.paragraph_format.space_after = Pt(10)

    doc.add_page_break()

    # --- 6. OUTPUT & DEMONSTRATION ---
    add_header()
    h6 = doc.add_paragraph()
    r6 = h6.add_run("6. OUTPUT & DEMONSTRATION:")
    r6.bold = True
    r6.font.size = Pt(14)
    r6.underline = True

    p_out_intro = doc.add_paragraph()
    p_out_intro.add_run("The AI Arithmetic Tutor system produces real-time interactive outputs across practice sessions, AI tutor chats, and PDF report downloads.")

    # Sample AI Output 1
    h_out1 = doc.add_paragraph()
    r_out1 = h_out1.add_run("Sample Output 1: AI Tutor Step-by-Step Mistake Breakdown")
    r_out1.bold = True
    r_out1.font.size = Pt(12)

    sample_ai_text = (
        "Input Problem: 14 + 8\n"
        "Student Answer: 20 (Incorrect)\n"
        "Correct Answer: 22\n\n"
        "AI Professor Owl Explanation Output:\n"
        "----------------------------------------------------------------------\n"
        "🦉 Professor Owl's Step-by-Step Breakdown:\n"
        "Great effort trying 14 + 8! Mistakes are just steps on the path to mastery.\n\n"
        "Step 1: Start with the First Number -> Imagine you have 14 blocks.\n"
        "Step 2: Add the Second Number -> Count forward 8 more: 14 + 1 + 1 + 1 + 1...\n"
        "Step 3: Combine Total -> Combining 14 and 8 gives exactly 22!\n\n"
        "Tip: You typed 20. You were very close! Check your ones column again.\n"
        "----------------------------------------------------------------------"
    )
    p_samp1 = doc.add_paragraph()
    r_s1 = p_samp1.add_run(sample_ai_text)
    r_s1.font.name = 'Courier New'
    r_s1.font.size = Pt(9.5)

    # Sample Output 2
    h_out2 = doc.add_paragraph()
    r_out2 = h_out2.add_run("Sample Output 2: PDF Academic Performance Report Export")
    r_out2.bold = True
    r_out2.font.size = Pt(12)

    sample_pdf_text = (
        "Generated Document: AI_Tutor_Report_DemoStudent.pdf\n"
        "Student Summary: DemoStudent | Grade Level: Grade 4 | XP: 240 ⭐ | Streak: 3 Days 🔥\n"
        "Operation Mastery:\n"
        "  - Addition: 20 Attempted, 18 Correct (90.0% Accuracy) [Level 3 ★★★☆☆]\n"
        "  - Subtraction: 15 Attempted, 12 Correct (80.0% Accuracy) [Level 2 ★★☆☆☆]\n"
        "  - Multiplication: 12 Attempted, 9 Correct (75.0% Accuracy) [Level 2 ★★☆☆☆]\n"
        "  - Division: 10 Attempted, 7 Correct (70.0% Accuracy) [Level 1 ★☆☆☆☆]\n"
        "Badges Earned: 🚀 First Step, 🔥 On Fire\n"
        "AI Recommendation: 'DemoStudent demonstrates solid consistency in fundamental addition. Recommended focus: Timed mixed division regrouping to achieve Level 3 mastery!'"
    )
    p_samp2 = doc.add_paragraph()
    r_s2 = p_samp2.add_run(sample_pdf_text)
    r_s2.font.name = 'Courier New'
    r_s2.font.size = Pt(9.5)

    # Save Document
    output_path = r"c:\Users\saras\OneDrive\Desktop\AI-Arithmetic-Tutor-Mobile-App\AI_Arithmetic_Tutor_College_Report.docx"
    doc.save(output_path)
    print("College Report DOCX created successfully at:", output_path)

if __name__ == "__main__":
    build_college_report()
